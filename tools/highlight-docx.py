#!/usr/bin/env python3
"""highlight-docx.py - colour the command blocks in the operating documents.

Usage:  highlight-docx.py <file.docx> [more.docx ...] [--apply]

Without --apply it reports what it would change and writes nothing.

WHY THIS EXISTS

The documents already colour their command blocks: command, --option,
$VARIABLE, "string" and # comment each get their own colour, and the manual
prints that legend in section 1.1. The colouring was applied by hand, and
hand-applied means it stops at whatever was written that day. Three
tr-inventory examples added after the fact arrived as a single purple run --
command, option and comment all one colour, which is worse than no colouring
because the legend says the colours mean something.

A blanket re-colour is the wrong fix: several blocks are directory trees and
sample output rather than commands, and their spacing and arrows were tuned
by hand. So this only touches a paragraph when the whole of its monospace
text is ONE colour and the text still parses as a command line. That is
precisely the signature of a block nobody highlighted, and it leaves
deliberate work alone.

THE PALETTE

Taken from the documents themselves, not invented here, so repaired blocks
are indistinguishable from the originals.
"""
import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    sys.exit("python-docx is missing. Run inside the kit venv:\n"
             "  ~/.translate-venv/bin/python tools/highlight-docx.py ...")

MONO = "Consolas"
SIZE = Pt(8.5)

COMMAND = RGBColor(0xA8, 0x43, 0x2A)   # bold
OPTION = RGBColor(0x1F, 0x6E, 0x8C)
VARIABLE = RGBColor(0x8A, 0x4B, 0x9E)
STRING = RGBColor(0x3F, 0x7A, 0x44)
COMMENT = RGBColor(0x8C, 0x85, 0x78)
PLAIN = RGBColor(0x35, 0x32, 0x2E)
OPERATOR = RGBColor(0xA3, 0x9C, 0x90)   # && || | ; and the -> in diagrams

# Commands the documents actually use. A first word not in this set is left
# plain rather than guessed at: colouring an output line's first word as a
# command is the mistake this tool exists to undo.
KIT = {"tr-project", "tr-inventory", "tr-status", "tr-run", "tr-lint",
       "tr-docx", "tr-txt", "tr-xlsx", "tr-pdf", "tr-ocrtext", "tr-setup",
       "tr-model", "tr-fixtures", "tr-hwsurvey", "case-init", "case-open",
       "case-close", "case-status", "case-guard", "case-guard-desktop",
       "ocr-check.py", "highlight-docx.py", "phase1-setup.sh",
       "install-desktop-guard.sh", "calibrate_lang.py", "cycle-test.sh"}
SHELL = {"cd", "ls", "cp", "mv", "rm", "mkdir", "chmod", "chattr", "sudo",
         "nano", "less", "cat", "grep", "git", "ollama", "python3", "bash",
         "sh", "xdg-open", "soffice", "tesseract", "pdftoppm", "pdftotext",
         "ocrmypdf", "df", "du", "free", "systemctl", "swapon", "export",
         "source", "echo", "wc", "sort", "head", "tail", "install", "sed",
         "claude", "tee", "mkswap", "fallocate", "cryptsetup", "mount",
         "umount", "mountpoint", "lsof", "screen", "tmux", "nohup", "watch"}
COMMANDS = KIT | SHELL

# A command may follow any of these, so the "first word" rule has to restart
# after them. Without this, `cd "$KIT" && claude` leaves claude plain, which
# is exactly the distinction the legend promises to show.
CHAIN = {"&&", "||", "|", ";", "sudo", "nohup", "watch", "then", "do"}

TOKEN = re.compile(r'''
    (?P<comment>\#.*$)
  | (?P<string>"[^"]*"|'[^']*')
  | (?P<op>&&|\|\||[|;])
  | (?P<option>(?<![\w-])--?[A-Za-z][\w-]*)
  | (?P<var>\$\{?[A-Za-z_][A-Za-z0-9_]*\}?|(?<![\w-])[A-Z_][A-Z0-9_]{1,}=)
  | (?P<space>\s+)
  | (?P<word>\S+)
''', re.VERBOSE)


def classify(text):
    """Split one command line into (text, colour, bold) pieces."""
    out = []
    first = True
    for m in TOKEN.finditer(text):
        kind = m.lastgroup
        tok = m.group()
        if kind == "comment":
            out.append((tok, COMMENT, False))
            first = False
        elif kind == "string":
            out.append((tok, STRING, False))
            first = False
        elif kind == "op":
            out.append((tok, OPERATOR, False))
            first = True          # a new command may start after it
        elif kind == "option":
            out.append((tok, OPTION, False))
            first = False
        elif kind == "var":
            out.append((tok, VARIABLE, False))
            first = False
        elif kind == "space":
            out.append((tok, PLAIN, False))
        else:
            base = os.path.basename(tok)
            if first and (tok in COMMANDS or base in COMMANDS):
                out.append((tok, COMMAND, True))
                first = tok in CHAIN
            else:
                out.append((tok, PLAIN, False))
                first = False
    return out


def looks_like_command(text):
    """A command line, not a tree diagram or a block of sample output."""
    t = text.strip()
    if not t or any(ch in t for ch in "┌┐└┘│─├┬>"):
        return False
    head = t.split()[0] if t.split() else ""
    if head.startswith("#"):
        return True
    return head in COMMANDS or os.path.basename(head) in COMMANDS


def mono_runs(p):
    return [r for r in p.runs if r.font.name and "Consol" in r.font.name]


def expected_colours(text):
    """One colour per character, as the palette rules would assign them."""
    out = []
    for tok, colour, _bold in classify(text):
        out.extend([str(colour)] * len(tok))
    return out


def actual_colours(p):
    """One colour per character, as the document currently has them."""
    out = []
    for r in p.runs:
        c = r.font.color.rgb if (r.font.color and r.font.color.rgb) else None
        out.extend([str(c) if c is not None else str(PLAIN)] * len(r.text))
    return out


def repaint(p):
    """Rewrite the paragraph's runs with the palette. Returns True if changed."""
    pieces = classify(p.text)
    if not pieces:
        return False
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for text, colour, bold in pieces:
        r = p.add_run(text)
        r.font.name = MONO
        r.font.size = SIZE
        r.font.color.rgb = colour
        r.bold = bold or None
    return True


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", nargs="+")
    ap.add_argument("--apply", action="store_true",
                    help="write the files; without it, only report")
    ap.add_argument("--strict", action="store_true",
                    help="also repaint blocks that merely differ from these "
                         "rules, not just ones painted a single colour")
    a = ap.parse_args()

    total = 0
    for path in a.docx:
        d = Document(path)
        hits = []
        for i, p in enumerate(d.paragraphs):
            if not mono_runs(p):
                continue
            if not looks_like_command(p.text):
                continue
            want = expected_colours(p.text)
            have = actual_colours(p)
            if have == want:
                continue
            # Default to repairing only what is actually broken: a line
            # carrying several token types that was painted a single colour.
            # That is the legend being contradicted, and it is what hand
            # editing leaves behind.
            #
            # Everything else that merely differs from these rules is left
            # alone unless --strict is given. Measured on the three
            # documents, --strict would rewrite 19 further paragraphs to
            # change `git clone` from one bold command into `git` bold plus
            # `clone` plain, and to recolour a heredoc's opening quote --
            # churn against hand-tuned blocks for no gain in clarity.
            if not a.strict and not (len(set(want)) > 1 and len(set(have)) == 1):
                continue
            hits.append((i, p))
        what = "mismatched" if a.strict else "unhighlighted"
        print(f"{os.path.basename(path)}: {len(hits)} {what} command block(s)")
        for i, p in hits:
            print(f"   p{i:<4} {p.text.strip()[:66]}")
            if a.apply:
                repaint(p)
        total += len(hits)
        if a.apply and hits:
            d.save(path)
            print(f"   written: {path}")
        print()

    if not a.apply and total:
        print("Nothing written. Re-run with --apply to fix.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
