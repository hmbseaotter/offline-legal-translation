#!/usr/bin/env python3
"""gen-docs.py - write the documented tables from lib/registry.py.

Usage:  gen-docs.py [--apply]      (default: check only, exit 1 on drift)

WHAT IT WRITES

  CLAUDE.md                    the Tools table and the environment table
  01-operating-manual.md       section 12's environment table
  01-operating-manual.docx     the same table, in the primary document

Each generated region is delimited by markers so the surrounding prose is
untouched:

    <!-- GENERATED:env --> ... <!-- /GENERATED:env -->

WHY THIS EXISTS

The same facts were written out in three places and drifted in all the ways
that allows. TR_PROMPT_VERSION was documented as v1 while the code had said
v4 for weeks. Six tools appeared in no document at all. Flags were added to
argparse and the usage lines were not updated. Every one of those read as
correct, because a table is not obviously stale the way a broken command is.

Testing that three copies agree finds the drift a commit late. Generating
two of them from the third means they cannot drift: there is one place to
edit, and forgetting to edit it is a different and much more visible kind of
mistake than editing one copy of three.

WHAT THE CHECK STILL CANNOT DO

It verifies that the documents match the registry, and that the registry's
defaults match the values actually written in the source. It cannot verify
that a purpose column is *true*. Prose stays a human problem; this removes
the mechanical half so the human half is all that is left.
"""
import argparse
import os
import re
import sys

KIT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.expanduser("~/translation-work/docs")
sys.path.insert(0, os.path.join(KIT, "lib"))
import registry  # noqa: E402


def md_table(rows, headers):
    widths = [max(len(str(r[i])) for r in [headers] + list(rows))
              for i in range(len(headers))]
    out = ["| " + " | ".join(h.ljust(widths[i]) for i, h in enumerate(headers)) + " |",
           "|" + "|".join("-" * (w + 2) for w in widths) + "|"]
    for r in rows:
        out.append("| " + " | ".join(str(c).ljust(widths[i])
                                     for i, c in enumerate(r)) + " |")
    return "\n".join(out)


def replace_region(text, tag, body):
    """Swap the contents between the markers for tag. Returns (text, found)."""
    open_m, close_m = f"<!-- GENERATED:{tag} -->", f"<!-- /GENERATED:{tag} -->"
    if open_m not in text or close_m not in text:
        return text, False
    i = text.index(open_m) + len(open_m)
    j = text.index(close_m)
    return text[:i] + "\n" + body + "\n" + text[j:], True


def env_rows():
    return [(n, d, p) for n, d, p in registry.ENV_VARS]


def code_defaults():
    """Defaults as the source actually sets them, for cross-checking."""
    found = {}
    pats = [
        re.compile(r'environ\.get\(\s*"([A-Z_]+)"\s*,\s*"([^"]*)"'),
        re.compile(r'\$\{([A-Z_]+):-([^}]*)\}'),
    ]
    for sub in ("bin", "lib", "tools"):
        d = os.path.join(KIT, sub)
        for f in sorted(os.listdir(d)):
            p = os.path.join(d, f)
            if not os.path.isfile(p):
                continue
            try:
                s = open(p, encoding="utf-8", errors="replace").read()
            except OSError:
                continue
            for pat in pats:
                for name, val in pat.findall(s):
                    found.setdefault(name, val)
    return found


def check_registry_against_code():
    """The registry claims defaults; the code sets them. Compare."""
    code = code_defaults()
    problems = []
    for name, default, _purpose in registry.ENV_VARS:
        if "/" in name:                      # "TR_SRC / TR_TGT" documents a pair
            continue
        if name not in code:
            continue                          # set elsewhere, or shell-only
        actual = code[name]
        # A default that is itself a variable -- ${TR_PROJECTS:-$default_projects}
        # in lib/guard.sh -- carries no literal to compare against. Skip it
        # rather than report a difference that is only about indirection.
        if actual.startswith("$"):
            continue
        want = default.strip()
        if want.startswith("~"):
            want = want.replace("~", "$HOME") if "$HOME" in actual else want
        if want in ("(empty)", "(unset)", "(active project)"):
            continue
        if actual != default and actual != want:
            problems.append(f"{name}: registry says {default!r}, "
                            f"code sets {actual!r}")
    return problems


def tool_reference():
    """Usage and flags for every command, read out of the tools themselves.

    Hand-written reference goes stale the moment a flag is added -- which is
    how --strict, --limit, --gate and --pin-all all came to exist in argparse
    and nowhere a user would look. Reading it from the source means the
    reference cannot disagree with the program.

    Python tools give their flags to argparse with help text, so those are
    parsed from the syntax tree. Shell tools have no such structure; their
    header comment is what there is, so the Usage line is taken verbatim and
    the flags live in it.
    """
    import ast
    out = []
    for sub in ("bin", "tools"):
        d = os.path.join(KIT, sub)
        for name in sorted(os.listdir(d)):
            path = os.path.join(d, name)
            if not os.path.isfile(path) or name.endswith((".tsv", ".pyc")):
                continue
            try:
                src = open(path, encoding="utf-8").read()
            except OSError:
                continue
            if not src.startswith("#!"):
                continue

            # One-line summary: "name - what it does", from the docstring or
            # the first comment line.
            summary = ""
            lines = src.splitlines()
            for i, ln in enumerate(lines):
                m = re.match(r'^(?:"""|#\s*)' + re.escape(name) + r'\s*[-–—]\s*(.+)$',
                             ln)
                if not m:
                    continue
                summary = m.group(1).strip().rstrip('"')
                # A header comment wraps. Without following it, case-guard was
                # described as "wrapper that refuses to launch Claude Code
                # while the case" -- cut mid-clause, which reads as a bug in
                # the tool rather than in the reference.
                is_py = src.startswith("#!/usr/bin/env python3")
                for cont in lines[i + 1:i + 4]:
                    if summary.endswith((".", "?", "!")):
                        break
                    raw = cont.strip()
                    # Only across the comment block itself. case-status ends
                    # its summary with "?", so without this the loop ran on
                    # into the script and documented the tool as
                    # "... safe to start Claude Code? set -uo pipefail".
                    if not (raw.startswith("#") or (is_py and '"""' not in raw)):
                        break
                    c = raw.lstrip("#").strip().rstrip('"')
                    if not c or c.startswith("Usage"):
                        break
                    summary += " " + c
                break

            usage = [l.strip().lstrip("#").strip()
                     for l in src.splitlines()
                     if re.match(r'^\s*#?\s*(Usage|usage):', l)]
            usage = [u for u in usage if u]

            flags = []
            if src.startswith("#!/usr/bin/env python3"):
                try:
                    tree = ast.parse(src)
                except SyntaxError:
                    tree = None
                for node in ast.walk(tree) if tree else []:
                    if (isinstance(node, ast.Call)
                            and isinstance(node.func, ast.Attribute)
                            and node.func.attr == "add_argument"):
                        names = [a.value for a in node.args
                                 if isinstance(a, ast.Constant)
                                 and isinstance(a.value, str)]
                        opt = [n for n in names if n.startswith("-")]
                        if not opt:
                            continue
                        help_txt = ""
                        for kw in node.keywords:
                            if kw.arg == "help" and isinstance(kw.value, ast.Constant):
                                help_txt = str(kw.value.value)
                            elif kw.arg == "help" and isinstance(kw.value, ast.JoinedStr):
                                help_txt = "".join(
                                    v.value for v in kw.value.values
                                    if isinstance(v, ast.Constant))
                        flags.append((" / ".join(opt), help_txt))
            out.append((f"{sub}/{name}" if sub == "tools" else name,
                        summary, usage, flags))
    return out


def reference_md():
    lines = []
    for name, summary, usage, flags in tool_reference():
        lines.append(f"### `{name}`")
        lines.append("")
        if summary:
            lines.append(summary)
            lines.append("")
        for u in usage:
            lines.append(f"    {u}")
        if usage:
            lines.append("")
        if flags:
            lines.append("| Flag | Meaning |")
            lines.append("|---|---|")
            for f, h in flags:
                lines.append(f"| `{f}` | {h or '—'} |")
            lines.append("")
    return "\n".join(lines).rstrip()


def write_docx_env(path, rows, apply):
    """Rewrite section 12's table in the primary document.

    Identified by its header cells rather than by index, so inserting a
    table earlier in the manual does not silently retarget this.
    """
    from docx import Document                     # kit venv only
    d = Document(path)
    target = None
    for t in d.tables:
        head = [c.text.strip() for c in t.rows[0].cells]
        if head[:3] == ["Variable", "Default", "Purpose"]:
            target = t
            break
    if target is None:
        return True                               # cannot find it: report drift

    current = [tuple(c.text.strip() for c in r.cells[:3])
               for r in target.rows[1:]]
    if current == [tuple(r) for r in rows]:
        return False
    if not apply:
        return True

    def set_cell(cell, text):
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.add_run(text)

    import copy
    # Clone the template from THIS table's own first data row, captured
    # before anything is deleted. Taking it from another table produced rows
    # with that table's column count -- two cells where three were needed --
    # and python-docx reports no error for it: the document simply comes out
    # with a truncated table.
    template = copy.deepcopy(target.rows[1]._tr) if len(target.rows) > 1 \
        else copy.deepcopy(target.rows[0]._tr)
    while len(target.rows) > 1:                    # keep the header
        target._tbl.remove(target.rows[-1]._tr)
    for row in rows:
        target._tbl.append(copy.deepcopy(template))
        cells = target.rows[-1].cells
        for cell, text in zip(cells, row):
            set_cell(cell, text)
    d.save(path)
    return False


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--apply", action="store_true",
                    help="write the files; without it, report drift only")
    a = ap.parse_args()

    drift = []

    problems = check_registry_against_code()
    for p in problems:
        drift.append(f"registry vs code: {p}")

    env_md = md_table(env_rows(), ("Variable", "Default", "Purpose"))
    tools_md = md_table(registry.TOOLS, ("Command", "Purpose"))

    targets = [
        (os.path.join(KIT, "CLAUDE.md"), [("env", env_md), ("tools", tools_md)]),
        (os.path.join(DOCS, "01-operating-manual.md"), [("env", env_md)]),
    ]

    # The runbook is written whole rather than patched between markers: its
    # prose half lives in docs-src/ and its reference half is read out of the
    # tools, so there is nothing in the output file for a person to edit and
    # nothing to preserve.
    intro = os.path.join(KIT, "docs-src", "runbook-intro.md")
    runbook = os.path.join(DOCS, "04-runbook.md")
    if os.path.exists(intro) and os.path.isdir(DOCS):
        body = open(intro, encoding="utf-8").read().rstrip() + "\n\n" \
            + reference_md() + "\n"
        cur = open(runbook, encoding="utf-8").read() if os.path.exists(runbook) else ""
        if cur != body:
            drift.append("04-runbook.md: out of date")
            if a.apply:
                open(runbook, "w", encoding="utf-8").write(body)

    for path, regions in targets:
        if not os.path.exists(path):
            # The kit can be cloned without the documents repository beside
            # it. That is a normal checkout, not drift, so skip rather than
            # block a commit on a file this clone was never expected to have.
            if path.startswith(DOCS):
                continue
            drift.append(f"missing: {path}")
            continue
        s = orig = open(path, encoding="utf-8").read()
        for tag, body in regions:
            s, found = replace_region(s, tag, body)
            if not found:
                drift.append(f"{os.path.basename(path)}: no GENERATED:{tag} markers")
        if s != orig:
            drift.append(f"{os.path.basename(path)}: table is out of date")
            if a.apply:
                open(path, "w", encoding="utf-8").write(s)

    # The .docx is the primary document and the .md is derived from it, so
    # leaving the docx hand-maintained would leave the authoritative copy as
    # the one that drifts. python-docx lives in the kit venv; without it,
    # say so rather than pass silently.
    docx = os.path.join(DOCS, "01-operating-manual.docx")
    if os.path.isdir(DOCS) and os.path.exists(docx):
        try:
            changed = write_docx_env(docx, env_rows(), a.apply)
            if changed:
                drift.append("01-operating-manual.docx: table is out of date")
        except ImportError:
            print("  note: python-docx unavailable, .docx table not checked.\n"
                  "        run under the kit venv: "
                  "~/.translate-venv/bin/python tools/gen-docs.py")

    if a.apply:
        print(f"  applied. {len(drift)} region(s) touched or flagged.")
        for d in drift:
            print(f"    {d}")
        return 1 if problems else 0

    if drift:
        print("Documented tables have drifted from lib/registry.py:")
        for d in drift:
            print(f"  {d}")
        print("\n  regenerate:  tools/gen-docs.py --apply")
        return 1
    print("  tables match lib/registry.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())
